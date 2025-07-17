"""
Document Extractor for CAHMS Neurodevelopmental Assessment
This module handles extraction of text content from various document formats.
"""

import logging
from typing import Dict, List, Optional, Any
from dataclasses import dataclass
from datetime import datetime

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


@dataclass
class DocumentContent:
    """Represents uploaded document content"""
    filename: str
    content: str
    document_type: str
    is_required: bool
    upload_timestamp: datetime


@dataclass
class AssessmentRequest:
    """Represents a complete assessment request"""
    documents: List[DocumentContent]
    session_id: str
    request_timestamp: datetime


class DocumentExtractor:
    """Handles extraction of text content from various document formats"""
    
    async def extract_document_content(self, file_content: bytes, filename: str) -> str:
        """
        Extract text content from uploaded document
        
        Args:
            file_content: Raw file content as bytes
            filename: Name of the uploaded file
            
        Returns:
            Extracted text content
        """
        try:
            # Check if file_content is valid
            if not file_content or len(file_content) == 0:
                logger.warning(f"Empty file content for {filename}")
                return f"[Empty file: {filename}]"
            
            # Log file size for debugging
            logger.info(f"Processing file {filename} with size {len(file_content)} bytes")
            
            # Extract based on file type
            if filename.lower().endswith('.txt'):
                # Handle text files
                try:
                    text = file_content.decode('utf-8')
                    return self._clean_text_content(text, filename)
                except UnicodeDecodeError:
                    try:
                        text = file_content.decode('latin-1', errors='ignore')
                        return self._clean_text_content(text, filename)
                    except Exception as e:
                        return f"[Error decoding text file {filename}: {str(e)}]"
                        
            elif filename.lower().endswith('.pdf'):
                return await self._extract_pdf_content(file_content, filename)
                
            elif filename.lower().endswith(('.doc', '.docx')):
                return await self._extract_docx_content(file_content, filename)
                
            else:
                logger.warning(f"Unsupported file format: {filename}")
                return f"[Unsupported file format: {filename}. Supported formats: .txt, .pdf, .doc, .docx]"
                
        except Exception as e:
            logger.error(f"Error extracting content from {filename}: {e}")
            return f"[Error extracting content from {filename}: {str(e)}]"
    
    def _clean_text_content(self, text: str, filename: str) -> str:
        """Clean and validate extracted text content"""
        import re
        
        # Remove null bytes and other problematic characters
        text = text.replace('\x00', '')
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        # Check if we have meaningful content
        if not text:
            return f"[File {filename} appears to be empty after processing]"
        
        # Check for suspiciously short content
        if len(text) < 10:
            return f"[File {filename} contains very little readable content: {text}]"
        
        # Check for binary data patterns (might indicate corruption)
        if len(re.findall(r'[^\x20-\x7E\n\r\t]', text)) > len(text) * 0.5:
            return f"[File {filename} appears to contain binary data or is corrupted]"
        
        return text
    
    async def _extract_pdf_content(self, file_content: bytes, filename: str) -> str:
        """Extract text from PDF files"""
        try:
            import PyPDF2
            import io
            
            # Check if file_content is not empty
            if not file_content or len(file_content) == 0:
                logger.warning(f"Empty file content for {filename}")
                return f"[Empty file: {filename}]"
            
            # Try to create PDF reader
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            
            # Check if PDF is encrypted
            if pdf_reader.is_encrypted:
                logger.warning(f"PDF file {filename} is encrypted")
                return f"[PDF file {filename} is encrypted and cannot be processed without a password]"
            
            # Extract text from all pages
            text = ""
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                except Exception as page_error:
                    logger.warning(f"Error extracting text from page {page_num + 1} of {filename}: {page_error}")
                    continue
            
            # Clean up extracted text
            text = text.strip()
            
            # Check if we got meaningful content
            if not text or len(text) < 10:
                logger.warning(f"Very little or no text extracted from {filename}")
                return f"[PDF file {filename} appears to contain no readable text content - it may be image-based or corrupted]"
            
            # Remove excessive whitespace and clean up
            import re
            text = re.sub(r'\s+', ' ', text)
            text = re.sub(r'\n\s*\n', '\n\n', text)
            
            return text
            
        except ImportError:
            logger.error(f"PyPDF2 not installed. Cannot process PDF file: {filename}")
            return f"[PDF content from {filename} - PyPDF2 not installed. Run: pip install PyPDF2]"
        except Exception as e:
            logger.error(f"Error extracting PDF content from {filename}: {e}")
            return f"[Error extracting PDF content from {filename}: {str(e)}]"
    
    async def _extract_docx_content(self, file_content: bytes, filename: str) -> str:
        """Extract text from DOCX files"""
        try:
            from docx import Document
            import io
            
            # Check if file_content is not empty
            if not file_content or len(file_content) == 0:
                logger.warning(f"Empty file content for {filename}")
                return f"[Empty file: {filename}]"
            
            # Check if it's a valid DOCX file
            try:
                doc = Document(io.BytesIO(file_content))
                text = ""
                
                # Extract text from paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text += paragraph.text + "\n"
                
                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            text += "\t".join(row_text) + "\n"
                
                # Clean up extracted text
                text = text.strip()
                
                # Check if we got meaningful content
                if not text or len(text) < 10:
                    logger.warning(f"Very little or no text extracted from {filename}")
                    return f"[DOCX file {filename} appears to contain no readable text content - it may be empty or corrupted]"
                
                # Remove excessive whitespace and clean up
                import re
                text = re.sub(r'\s+', ' ', text)
                text = re.sub(r'\n\s*\n', '\n\n', text)
                
                return text
                
            except Exception as docx_error:
                # If DOCX parsing fails, try to extract as plain text
                logger.warning(f"DOCX parsing failed for {filename}, trying text extraction: {docx_error}")
                
                # Check if it might be a password-protected file
                if "password" in str(docx_error).lower() or "encrypted" in str(docx_error).lower():
                    return f"[DOCX file {filename} appears to be password-protected or encrypted]"
                
                try:
                    # Try to decode as UTF-8 text
                    text_content = file_content.decode('utf-8', errors='ignore')
                    
                    # Clean up and check if we got meaningful content
                    text_content = text_content.strip()
                    if len(text_content) > 50:  # Arbitrary threshold for meaningful content
                        return text_content
                    else:
                        # If that fails, try latin-1 encoding
                        try:
                            text_content = file_content.decode('latin-1', errors='ignore')
                            text_content = text_content.strip()
                            if len(text_content) > 50:
                                return text_content
                            else:
                                return f"[Could not extract meaningful content from {filename} - file may be corrupted, empty, or in unsupported format]"
                        except:
                            return f"[Could not extract content from {filename} - file may be corrupted or in unsupported format]"
                except:
                    return f"[Could not extract content from {filename} - file may be corrupted or in unsupported format]"
                        
        except ImportError:
            logger.error(f"python-docx not installed. Cannot process DOCX file: {filename}")
            return f"[DOCX content from {filename} - python-docx not installed. Run: pip install python-docx]"
        except Exception as e:
            logger.error(f"Error extracting DOCX content from {filename}: {e}")
            return f"[Error extracting DOCX content from {filename}: {str(e)}]"
    
    async def validate_documents(self, documents: List[DocumentContent]) -> Dict[str, Any]:
        """
        Validate that uploaded documents contain appropriate content
        
        Args:
            documents: List of documents to validate
            
        Returns:
            Validation results
        """
        validation_results = {
            "valid": True,
            "warnings": [],
            "errors": [],
            "document_analysis": {}
        }
        
        # Check for required documents based on what's actually marked as required
        required_docs = [doc for doc in documents if doc.is_required]
        
        # Check if any required documents are missing (empty or no content)
        missing_required = []
        for doc in required_docs:
            if not doc.content or len(doc.content.strip()) == 0:
                missing_required.append(doc.document_type)
        
        if missing_required:
            validation_results["valid"] = False
            validation_results["errors"].append(f"Required documents have no content: {', '.join(missing_required)}")
        
        # Analyze each document
        for doc in documents:
            analysis = {
                "word_count": len(doc.content.split()),
                "has_content": len(doc.content.strip()) > 0,
                "file_size_kb": len(doc.content.encode('utf-8')) / 1024
            }
            
            if not analysis["has_content"]:
                if doc.is_required:
                    validation_results["errors"].append(f"{doc.document_type} is required but appears to be empty")
                else:
                    validation_results["warnings"].append(f"{doc.document_type} appears to be empty")
            elif analysis["word_count"] < 10:
                validation_results["warnings"].append(f"{doc.document_type} has very little content ({analysis['word_count']} words)")
            
            validation_results["document_analysis"][doc.document_type] = analysis
        
        return validation_results


async def process_assessment_documents(uploaded_files: Dict[str, Any], session_id: str, mandatory_uploads: Dict[str, bool] = None) -> AssessmentRequest:
    """
    Process uploaded files into an AssessmentRequest
    
    Args:
        uploaded_files: Dictionary of uploaded files from Streamlit session state
        session_id: Unique session identifier
        mandatory_uploads: Dictionary specifying which uploads are mandatory
        
    Returns:
        AssessmentRequest object
    """
    documents = []
    
    # Default mandatory configuration if not provided
    if mandatory_uploads is None:
        mandatory_uploads = {
            "form_s": True,
            "form_h": False,
            "form_a": False,
            "cahms_initial": False,
            "neuro_dev_history": False,
            "formulation_document": False,
            "school_observation": False,
            "supporting_information": False
        }
    
    # Define document mappings
    document_mappings = {
        "form_s": "Form S",
        "form_h": "Form H",
        "form_a": "Form A",
        "cahms_initial": "CAHMS Initial Assessment",
        "neuro_dev_history": "Neuro Dev History",
        "formulation_document": "Formulation Document",
        "school_observation": "School Observation",
        "supporting_information": "Supporting Information"
    }
    
    # Initialize document extractor
    extractor = DocumentExtractor()
    
    for file_key, doc_type in document_mappings.items():
        file_obj = uploaded_files.get(file_key)
        is_required = mandatory_uploads.get(file_key, False)
        
        if file_obj:
            try:
                # Read file content
                file_content = file_obj.read()
                
                # Extract text content using the document extractor
                text_content = await extractor.extract_document_content(file_content, file_obj.name)
                
                # Create document content object
                doc_content = DocumentContent(
                    filename=file_obj.name,
                    content=text_content,
                    document_type=doc_type,
                    is_required=is_required,
                    upload_timestamp=datetime.now()
                )
                
                documents.append(doc_content)
                logger.info(f"Successfully processed {doc_type} ({file_obj.name})")
                
            except Exception as e:
                logger.error(f"Error processing {doc_type} ({file_obj.name}): {e}")
                # Create a placeholder document with error message
                doc_content = DocumentContent(
                    filename=file_obj.name,
                    content=f"[Error processing {file_obj.name}: {str(e)}]",
                    document_type=doc_type,
                    is_required=is_required,
                    upload_timestamp=datetime.now()
                )
                documents.append(doc_content)
    
    return AssessmentRequest(
        documents=documents,
        session_id=session_id,
        request_timestamp=datetime.now()
    )
