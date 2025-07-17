"""
Azure LLM Client for CAHMS Neurodevelopmental Assessment
This module handles interactions with Azure-hosted Large Language Models
for processing uploaded documents and generating assessment reports.
"""

import os
import json
import logging
from typing import Dict, List, Optional, Any
from dataclasses import dataclass
import asyncio
from datetime import datetime

# LLM API imports
try:
    import requests
    from openai import OpenAI
except ImportError:
    print("Required packages not installed. Install with: pip install openai requests")
    OpenAI = None
    requests = None

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


@dataclass
class DocumentContent:
    """Represents uploaded document content"""
    filename: str
    content: str
    document_type: str
    is_required: bool
    upload_timestamp: datetime


@dataclass
class AssessmentRequest:
    """Represents a complete assessment request"""
    documents: List[DocumentContent]
    session_id: str
    request_timestamp: datetime


class AzureLLMClient:
    """Client for interacting with LLM services via API endpoint"""
    
    def __init__(self):
        """Initialize the LLM client with configuration from environment variables"""
        self.endpoint = os.getenv("LLM_ENDPOINT")
        self.api_key = os.getenv("LLM_API_KEY")
        self.model_name = os.getenv("LLM_MODEL_NAME", "gpt-4")
        self.use_openai_client = os.getenv("USE_OPENAI_CLIENT", "true").lower() == "true"
        
        # Initialize client
        self.client = None
        self._initialize_client()
    
    def _initialize_client(self):
        """Initialize the LLM client"""
        if not self.endpoint or not self.api_key:
            logger.warning("LLM credentials not configured. Set LLM_ENDPOINT and LLM_API_KEY")
            return
        
        try:
            if self.use_openai_client and OpenAI:
                # Use OpenAI client with Azure OpenAI configuration
                from openai import AzureOpenAI
                self.client = AzureOpenAI(
                    azure_endpoint=self.endpoint,
                    api_key=self.api_key,
                    api_version="2024-02-15-preview"
                )
                logger.info("Azure OpenAI client initialized successfully")
            else:
                # Use direct HTTP requests
                logger.info("Using direct HTTP requests for LLM communication")
                
        except ImportError:
            logger.warning("AzureOpenAI not available, falling back to standard OpenAI client")
            try:
                # Fallback to standard OpenAI client
                self.client = OpenAI(
                    base_url=f"{self.endpoint.rstrip('/')}/openai/deployments/{self.model_name}",
                    api_key=self.api_key,
                    default_headers={"api-version": "2024-02-15-preview"}
                )
                logger.info("Standard OpenAI client initialized with Azure endpoint")
            except Exception as e:
                logger.error(f"Failed to initialize OpenAI client: {e}")
        except Exception as e:
            logger.error(f"Failed to initialize LLM client: {e}")
    
    def is_configured(self) -> bool:
        """Check if the client is properly configured"""
        return (self.endpoint is not None and self.api_key is not None)
    
    def _make_direct_api_call(self, messages: List[Dict[str, str]], **kwargs) -> Dict[str, Any]:
        """
        Make a direct API call to the LLM endpoint
        
        Args:
            messages: List of messages for the conversation
            **kwargs: Additional parameters for the API call
            
        Returns:
            API response as dictionary
        """
        if not requests:
            raise ImportError("requests library not available")
        
        headers = {
            "Content-Type": "application/json",
            "api-key": self.api_key,  # Azure OpenAI uses 'api-key' header
            "api-version": "2024-02-15-preview"
        }
        
        payload = {
            "messages": messages,
            "max_tokens": kwargs.get("max_tokens", 4000),
            "temperature": kwargs.get("temperature", 0.3),
            "top_p": kwargs.get("top_p", 0.9)
        }
        
        # Construct the correct Azure OpenAI URL
        url = f"{self.endpoint.rstrip('/')}/openai/deployments/{self.model_name}/chat/completions"
        
        try:
            response = requests.post(
                url,
                headers=headers,
                json=payload,
                timeout=120
            )
            response.raise_for_status()
            return response.json()
        except Exception as e:
            logger.error(f"Direct API call failed: {e}")
            raise
    
    async def extract_document_content(self, file_content: bytes, filename: str) -> str:
        """
        Extract text content from uploaded document
        
        Args:
            file_content: Raw file content as bytes
            filename: Name of the uploaded file
            
        Returns:
            Extracted text content
        """
        try:
            if filename.lower().endswith('.txt'):
                return file_content.decode('utf-8')
            elif filename.lower().endswith('.pdf'):
                try:
                    import PyPDF2
                    import io
                    pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                    text = ""
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                    return text.strip()
                except ImportError:
                    logger.error(f"PyPDF2 not installed. Cannot process PDF file: {filename}")
                    return f"[PDF content from {filename} - PyPDF2 not installed. Run: pip install PyPDF2]"
                except Exception as e:
                    logger.error(f"Error extracting PDF content from {filename}: {e}")
                    return f"[Error extracting PDF content from {filename}: {str(e)}]"
            elif filename.lower().endswith(('.doc', '.docx')):
                try:
                    from docx import Document
                    import io
                    
                    # Check if it's a valid DOCX file
                    try:
                        doc = Document(io.BytesIO(file_content))
                        text = ""
                        for paragraph in doc.paragraphs:
                            text += paragraph.text + "\n"
                        # Also extract text from tables
                        for table in doc.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    text += cell.text + "\t"
                                text += "\n"
                        return text.strip()
                    except Exception as docx_error:
                        # If DOCX parsing fails, try to extract as plain text
                        logger.warning(f"DOCX parsing failed for {filename}, trying text extraction: {docx_error}")
                        try:
                            # Try to decode as UTF-8 text
                            text_content = file_content.decode('utf-8', errors='ignore')
                            return text_content.strip()
                        except:
                            # If that fails, try latin-1 encoding
                            try:
                                text_content = file_content.decode('latin-1', errors='ignore')
                                return text_content.strip()
                            except:
                                return f"[Could not extract content from {filename} - file may be corrupted or in unsupported format]"
                                
                except ImportError:
                    logger.error(f"python-docx not installed. Cannot process DOCX file: {filename}")
                    return f"[DOCX content from {filename} - python-docx not installed. Run: pip install python-docx]"
                except Exception as e:
                    logger.error(f"Error extracting DOCX content from {filename}: {e}")
                    return f"[Error extracting DOCX content from {filename}: {str(e)}]"
            else:
                return f"[Unsupported file format: {filename}]"
        except Exception as e:
            logger.error(f"Error extracting content from {filename}: {e}")
            return f"[Error extracting content from {filename}: {str(e)}]"
    
    def _create_assessment_prompt(self, documents: List[DocumentContent]) -> str:
        """
        Create a comprehensive prompt for the LLM based on uploaded documents
        
        Args:
            documents: List of processed documents
            
        Returns:
            Formatted prompt for the LLM
        """
        prompt = """You are a specialist clinician working with the CAMHS (Child and Adolescent Mental Health Services) Neurodevelopmental Team. 

Your task is to analyze the provided assessment documents and generate a comprehensive neurodevelopmental assessment report.

AVAILABLE DOCUMENTS:
"""
        
        for doc in documents:
            status = "REQUIRED" if doc.is_required else "OPTIONAL"
            prompt += f"\n{status} - {doc.document_type} ({doc.filename}):\n"
            prompt += f"{doc.content}\n"
            prompt += "-" * 80 + "\n"
        
        prompt += """
ASSESSMENT TASK:
Please analyze all the provided documents and generate a comprehensive neurodevelopmental assessment report that includes:

1. **Executive Summary**
   - Key findings and recommendations
   - Primary neurodevelopmental concerns identified

2. **Background Information**
   - Relevant developmental history
   - Family context and environmental factors
   - Educational background and challenges

3. **Assessment Findings**
   - Analysis of attention, concentration, and hyperactivity (Form H)
   - Social interaction and communication assessment (Form A)
   - School-based observations and support needs (Form S)
   - Clinical observations from initial assessment

4. **Formulation**
   - Integration of all assessment information
   - Identified strengths and challenges
   - Neurodevelopmental profile

5. **Recommendations**
   - Intervention strategies
   - Support recommendations for school and home
   - Follow-up requirements
   - Referrals to other services if needed

6. **Conclusion**
   - Summary of key points
   - Prognosis and next steps

Please ensure the report is:
- Professional and clinical in tone
- Evidence-based using the provided documentation
- Comprehensive yet accessible
- Structured with clear headings and bullet points where appropriate
- Sensitive to the young person and family's needs

Generate the assessment report now:
"""
        return prompt
    
    async def generate_assessment_report(self, assessment_request: AssessmentRequest) -> Dict[str, Any]:
        """
        Generate a neurodevelopmental assessment report using LLM
        
        Args:
            assessment_request: Complete assessment request with documents
            
        Returns:
            Dictionary containing the generated report and metadata
        """
        if not self.is_configured():
            return {
                "success": False,
                "error": "LLM client not configured properly",
                "report": None
            }
        
        try:
            # Create the assessment prompt
            prompt = self._create_assessment_prompt(assessment_request.documents)
            
            messages = [
                {
                    "role": "system", 
                    "content": "You are an expert clinician specializing in neurodevelopmental assessments for children and adolescents."
                },
                {
                    "role": "user", 
                    "content": prompt
                }
            ]
            
            # Call LLM - try OpenAI client first, then direct API
            if self.client and self.use_openai_client:
                response = self.client.chat.completions.create(
                    model=self.model_name,  # Model is still required for Azure OpenAI
                    messages=messages,
                    max_tokens=4000,
                    temperature=0.3,  # Lower temperature for more consistent clinical reports
                    top_p=0.9
                )
                report_content = response.choices[0].message.content
                tokens_used = response.usage.total_tokens if hasattr(response, 'usage') else None
            else:
                # Use direct API call
                response = self._make_direct_api_call(
                    messages=messages,
                    max_tokens=4000,
                    temperature=0.3,
                    top_p=0.9
                )
                report_content = response['choices'][0]['message']['content']
                tokens_used = response.get('usage', {}).get('total_tokens')
            
            # Create response with metadata
            result = {
                "success": True,
                "report": report_content,
                "metadata": {
                    "session_id": assessment_request.session_id,
                    "generation_timestamp": datetime.now().isoformat(),
                    "model_used": self.model_name,
                    "documents_processed": len(assessment_request.documents),
                    "required_documents": len([d for d in assessment_request.documents if d.is_required]),
                    "optional_documents": len([d for d in assessment_request.documents if not d.is_required]),
                    "tokens_used": tokens_used
                },
                "error": None
            }
            
            logger.info(f"Successfully generated assessment report for session {assessment_request.session_id}")
            return result
            
        except Exception as e:
            logger.error(f"Error generating assessment report: {e}")
            return {
                "success": False,
                "error": str(e),
                "report": None,
                "metadata": {
                    "session_id": assessment_request.session_id,
                    "generation_timestamp": datetime.now().isoformat(),
                    "error_timestamp": datetime.now().isoformat()
                }
            }
    
    async def validate_documents(self, documents: List[DocumentContent]) -> Dict[str, Any]:
        """
        Validate that uploaded documents contain appropriate content
        
        Args:
            documents: List of documents to validate
            
        Returns:
            Validation results
        """
        validation_results = {
            "valid": True,
            "warnings": [],
            "errors": [],
            "document_analysis": {}
        }
        
        # Check for required documents based on what's actually marked as required
        required_docs = [doc for doc in documents if doc.is_required]
        uploaded_required_types = {doc.document_type for doc in required_docs}
        
        # Check if any required documents are missing (empty or no content)
        missing_required = []
        for doc in required_docs:
            if not doc.content or len(doc.content.strip()) == 0:
                missing_required.append(doc.document_type)
        
        if missing_required:
            validation_results["valid"] = False
            validation_results["errors"].append(f"Required documents have no content: {', '.join(missing_required)}")
        
        # Analyze each document
        for doc in documents:
            analysis = {
                "word_count": len(doc.content.split()),
                "has_content": len(doc.content.strip()) > 0,
                "file_size_kb": len(doc.content.encode('utf-8')) / 1024
            }
            
            if not analysis["has_content"]:
                if doc.is_required:
                    validation_results["errors"].append(f"{doc.document_type} is required but appears to be empty")
                else:
                    validation_results["warnings"].append(f"{doc.document_type} appears to be empty")
            elif analysis["word_count"] < 10:
                validation_results["warnings"].append(f"{doc.document_type} has very little content ({analysis['word_count']} words)")
            
            validation_results["document_analysis"][doc.document_type] = analysis
        
        return validation_results


# Utility functions for integration with Streamlit app

def create_document_content(file, document_type: str, is_required: bool = True) -> DocumentContent:
    """
    Helper function to create DocumentContent from Streamlit file upload
    
    Args:
        file: Streamlit uploaded file object
        document_type: Type of document
        is_required: Whether the document is required
        
    Returns:
        DocumentContent object
    """
    if file is None:
        return None
    
    try:
        # Read file content
        content = file.read()
        if isinstance(content, bytes):
            content = content.decode('utf-8')
        
        return DocumentContent(
            filename=file.name,
            content=content,
            document_type=document_type,
            is_required=is_required,
            upload_timestamp=datetime.now()
        )
    except Exception as e:
        logger.error(f"Error creating document content for {file.name}: {e}")
        return None


async def process_assessment_documents(uploaded_files: Dict[str, Any], session_id: str, mandatory_uploads: Dict[str, bool] = None) -> AssessmentRequest:
    """
    Process uploaded files into an AssessmentRequest
    
    Args:
        uploaded_files: Dictionary of uploaded files from Streamlit session state
        session_id: Unique session identifier
        mandatory_uploads: Dictionary specifying which uploads are mandatory
        
    Returns:
        AssessmentRequest object
    """
    documents = []
    
    # Default mandatory configuration if not provided
    if mandatory_uploads is None:
        mandatory_uploads = {
            "form_s": True,
            "form_h": False,
            "form_a": False,
            "cahms_initial": False,
            "neuro_dev_history": False,
            "formulation_document": False,
            "school_observation": False,
            "supporting_information": False
        }
    
    # Define document mappings
    document_mappings = {
        "form_s": "Form S",
        "form_h": "Form H",
        "form_a": "Form A",
        "cahms_initial": "CAHMS Initial Assessment",
        "neuro_dev_history": "Neuro Dev History",
        "formulation_document": "Formulation Document",
        "school_observation": "School Observation",
        "supporting_information": "Supporting Information"
    }
    
    # Initialize LLM client for document processing
    llm_client = AzureLLMClient()
    
    for file_key, doc_type in document_mappings.items():
        file_obj = uploaded_files.get(file_key)
        is_required = mandatory_uploads.get(file_key, False)
        
        if file_obj:
            try:
                # Read file content
                file_content = file_obj.read()
                
                # Extract text content using the LLM client
                text_content = await llm_client.extract_document_content(file_content, file_obj.name)
                
                # Create document content object
                doc_content = DocumentContent(
                    filename=file_obj.name,
                    content=text_content,
                    document_type=doc_type,
                    is_required=is_required,
                    upload_timestamp=datetime.now()
                )
                
                documents.append(doc_content)
                logger.info(f"Successfully processed {doc_type} ({file_obj.name})")
                
            except Exception as e:
                logger.error(f"Error processing {doc_type} ({file_obj.name}): {e}")
                # Create a placeholder document with error message
                doc_content = DocumentContent(
                    filename=file_obj.name,
                    content=f"[Error processing {file_obj.name}: {str(e)}]",
                    document_type=doc_type,
                    is_required=is_required,
                    upload_timestamp=datetime.now()
                )
                documents.append(doc_content)
    
    return AssessmentRequest(
        documents=documents,
        session_id=session_id,
        request_timestamp=datetime.now()
    )