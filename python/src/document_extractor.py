"""
Generic Document Extractor
This module handles extraction of text content from various document formats.
Supports PDF, DOCX, and text files with robust error handling and validation.
"""

import logging
import re
import io
from typing import Dict, List, Optional, Any
from dataclasses import dataclass
from datetime import datetime

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


@dataclass
class DocumentContent:
    """Represents uploaded document content"""
    filename: str
    content: str
    document_type: str
    is_required: bool
    upload_timestamp: datetime


@dataclass
class DocumentRequest:
    """Represents a complete document processing request"""
    documents: List[DocumentContent]
    session_id: str
    request_timestamp: datetime


class DocumentExtractor:
    """Handles extraction of text content from various document formats"""
    
    async def extract_document_content(self, file_content: bytes, filename: str) -> str:
        """
        Extract text content from uploaded document
        
        Args:
            file_content: Raw file content as bytes
            filename: Name of the uploaded file
            
        Returns:
            Extracted text content
        """
        try:
            if not file_content:
                logger.warning(f"Empty file content for {filename}")
                return f"[Empty file: {filename}]"
            
            logger.info(f"Processing file {filename} with size {len(file_content)} bytes")
            
            # Extract based on file type
            if filename.lower().endswith('.txt'):
                return self._extract_text_content(file_content, filename)
            elif filename.lower().endswith('.pdf'):
                return await self._extract_pdf_content(file_content, filename)
            elif filename.lower().endswith(('.doc', '.docx')):
                return await self._extract_docx_content(file_content, filename)
            else:
                logger.warning(f"Unsupported file format: {filename}")
                return f"[Unsupported file format: {filename}. Supported formats: .txt, .pdf, .doc, .docx]"
                
        except Exception as e:
            logger.error(f"Error extracting content from {filename}: {e}")
            return f"[Error extracting content from {filename}: {str(e)}]"
    
    def _extract_text_content(self, file_content: bytes, filename: str) -> str:
        """Extract and clean text content from text files"""
        try:
            text = file_content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                text = file_content.decode('latin-1', errors='ignore')
            except Exception as e:
                return f"[Error decoding text file {filename}: {str(e)}]"
        
        return self._clean_text_content(text, filename)
    
    def _clean_text_content(self, text: str, filename: str) -> str:
        """Clean and validate extracted text content"""
        # Remove null bytes and problematic characters
        text = text.replace('\x00', '')
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = text.strip()
        
        # Validate content
        if not text:
            return f"[File {filename} appears to be empty after processing]"
        
        if len(text) < 10:
            return f"[File {filename} contains very little readable content: {text}]"
        
        # Check for binary data patterns
        if len(re.findall(r'[^\x20-\x7E\n\r\t]', text)) > len(text) * 0.5:
            return f"[File {filename} appears to contain binary data or is corrupted]"
        
        return text
    
    async def _extract_pdf_content(self, file_content: bytes, filename: str) -> str:
        """Extract text from PDF files"""
        try:
            import PyPDF2
            
            if not file_content:
                return f"[Empty file: {filename}]"
            
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            
            if pdf_reader.is_encrypted:
                logger.warning(f"PDF file {filename} is encrypted")
                return f"[PDF file {filename} is encrypted and cannot be processed without a password]"
            
            # Extract text from all pages
            text = ""
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                except Exception as page_error:
                    logger.warning(f"Error extracting text from page {page_num + 1} of {filename}: {page_error}")
                    continue
            
            text = text.strip()
            
            if not text or len(text) < 10:
                logger.warning(f"Very little or no text extracted from {filename}")
                return f"[PDF file {filename} appears to contain no readable text content - it may be image-based or corrupted]"
            
            # Clean up extracted text
            text = re.sub(r'\s+', ' ', text)
            text = re.sub(r'\n\s*\n', '\n\n', text)
            
            return text
            
        except ImportError:
            logger.error(f"PyPDF2 not installed. Cannot process PDF file: {filename}")
            return f"[PDF content from {filename} - PyPDF2 not installed. Run: pip install PyPDF2]"
        except Exception as e:
            logger.error(f"Error extracting PDF content from {filename}: {e}")
            return f"[Error extracting PDF content from {filename}: {str(e)}]"
    
    async def _extract_docx_content(self, file_content: bytes, filename: str) -> str:
        """Extract text from DOCX files"""
        try:
            from docx import Document
            
            if not file_content:
                return f"[Empty file: {filename}]"
            
            try:
                doc = Document(io.BytesIO(file_content))
                text = ""
                
                # Extract text from paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text += paragraph.text + "\n"
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:
                            text += "\t".join(row_text) + "\n"
                
                text = text.strip()
                
                if not text or len(text) < 10:
                    logger.warning(f"Very little or no text extracted from {filename}")
                    return f"[DOCX file {filename} appears to contain no readable text content - it may be empty or corrupted]"
                
                # Clean up extracted text
                text = re.sub(r'\s+', ' ', text)
                text = re.sub(r'\n\s*\n', '\n\n', text)
                
                return text
                
            except Exception as docx_error:
                logger.warning(f"DOCX parsing failed for {filename}, trying text extraction: {docx_error}")
                
                if "password" in str(docx_error).lower() or "encrypted" in str(docx_error).lower():
                    return f"[DOCX file {filename} appears to be password-protected or encrypted]"
                
                # Fallback to plain text extraction
                for encoding in ['utf-8', 'latin-1']:
                    try:
                        text_content = file_content.decode(encoding, errors='ignore').strip()
                        if len(text_content) > 50:
                            return text_content
                    except:
                        continue
                
                return f"[Could not extract content from {filename} - file may be corrupted or in unsupported format]"
                        
        except ImportError:
            logger.error(f"python-docx not installed. Cannot process DOCX file: {filename}")
            return f"[DOCX content from {filename} - python-docx not installed. Run: pip install python-docx]"
        except Exception as e:
            logger.error(f"Error extracting DOCX content from {filename}: {e}")
            return f"[Error extracting DOCX content from {filename}: {str(e)}]"
    
    async def preprocess_documents(self, documents: List[DocumentContent]) -> List[DocumentContent]:
        """
        Preprocess documents to check for issues and ensure they're in the correct format
        
        Args:
            documents: List of DocumentContent objects to preprocess
            
        Returns:
            List of processed DocumentContent objects
        """
        logger.info(f"Preprocessing {len(documents)} documents")
        processed_documents = []
        
        for i, doc in enumerate(documents):
            logger.debug(f"Processing document {i+1}: {doc.filename}")
            
            try:
                # Check if document has processing errors
                if doc.content.startswith('[') and doc.content.endswith(']'):
                    logger.warning(f"Document {doc.filename} has processing error: {doc.content}")
                else:
                    # Check for suspiciously short content
                    if len(doc.content.strip()) < 50:
                        logger.warning(f"Document {doc.filename} has very short content: {len(doc.content)} characters")
                    
                    # Check for binary data patterns
                    non_printable_ratio = len(re.findall(r'[^\x20-\x7E\n\r\t]', doc.content)) / len(doc.content) if doc.content else 0
                    if non_printable_ratio > 0.3:
                        logger.warning(f"Document {doc.filename} contains high ratio of non-printable characters: {non_printable_ratio:.2%}")
                    
                    logger.debug(f"Document {i+1} processed successfully ({len(doc.content)} chars)")
                
                processed_documents.append(doc)
                    
            except Exception as e:
                logger.error(f"Error processing document {doc.filename}: {e}")
                processed_documents.append(doc)
        
        logger.info(f"Preprocessing complete, {len(processed_documents)} documents ready")
        return processed_documents
    
    async def validate_documents(self, documents: List[DocumentContent]) -> Dict[str, Any]:
        """
        Validate that uploaded documents contain appropriate content
        
        Args:
            documents: List of documents to validate
            
        Returns:
            Validation results
        """
        validation_results = {
            "valid": True,
            "warnings": [],
            "errors": [],
            "document_analysis": {}
        }
        
        # Check for required documents that are missing content
        required_docs = [doc for doc in documents if doc.is_required]
        missing_required = [doc.document_type for doc in required_docs if not doc.content or not doc.content.strip()]
        
        if missing_required:
            validation_results["valid"] = False
            validation_results["errors"].append(f"Required documents have no content: {', '.join(missing_required)}")
        
        # Analyze each document
        for doc in documents:
            analysis = {
                "word_count": len(doc.content.split()),
                "has_content": len(doc.content.strip()) > 0,
                "file_size_kb": len(doc.content.encode('utf-8')) / 1024
            }
            
            if not analysis["has_content"]:
                if doc.is_required:
                    validation_results["errors"].append(f"{doc.document_type} is required but appears to be empty")
                else:
                    validation_results["warnings"].append(f"{doc.document_type} appears to be empty")
            elif analysis["word_count"] < 10:
                validation_results["warnings"].append(f"{doc.document_type} has very little content ({analysis['word_count']} words)")
            
            validation_results["document_analysis"][doc.document_type] = analysis
        
        return validation_results


async def process_documents(
    uploaded_files: Dict[str, Any], 
    session_id: str, 
    document_mappings: Dict[str, str] = None,
    mandatory_uploads: Dict[str, bool] = None
) -> DocumentRequest:
    """
    Process uploaded files into a DocumentRequest
    
    Args:
        uploaded_files: Dictionary of uploaded files from file upload interface
        session_id: Unique session identifier
        document_mappings: Dictionary mapping file keys to document type names
        mandatory_uploads: Dictionary specifying which uploads are mandatory
        
    Returns:
        DocumentRequest object
    """
    documents = []
    
    # Default document mappings if not provided
    if document_mappings is None:
        document_mappings = {
            "document_1": "Document 1",
            "document_2": "Document 2",
            "document_3": "Document 3"
        }
    
    if mandatory_uploads is None:
        mandatory_uploads = {key: False for key in document_mappings.keys()}
    
    extractor = DocumentExtractor()
    
    for file_key, doc_type in document_mappings.items():
        file_obj = uploaded_files.get(file_key)
        is_required = mandatory_uploads.get(file_key, False)
        
        if file_obj:
            try:
                file_content = file_obj.read()
                text_content = await extractor.extract_document_content(file_content, file_obj.name)
                
                doc_content = DocumentContent(
                    filename=file_obj.name,
                    content=text_content,
                    document_type=doc_type,
                    is_required=is_required,
                    upload_timestamp=datetime.now()
                )
                
                documents.append(doc_content)
                logger.info(f"Successfully processed {doc_type} ({file_obj.name})")
                
            except Exception as e:
                logger.error(f"Error processing {doc_type} ({file_obj.name}): {e}")
                doc_content = DocumentContent(
                    filename=file_obj.name,
                    content=f"[Error processing {file_obj.name}: {str(e)}]",
                    document_type=doc_type,
                    is_required=is_required,
                    upload_timestamp=datetime.now()
                )
                documents.append(doc_content)
    
    return DocumentRequest(
        documents=documents,
        session_id=session_id,
        request_timestamp=datetime.now()
    )


# Convenience functions for specific use cases

async def process_assessment_documents(uploaded_files: Dict[str, Any], session_id: str, mandatory_uploads: Dict[str, bool] = None) -> DocumentRequest:
    """
    Process CAHMS assessment documents - convenience wrapper for process_documents
    
    Args:
        uploaded_files: Dictionary of uploaded files from Streamlit session state
        session_id: Unique session identifier
        mandatory_uploads: Dictionary specifying which uploads are mandatory
        
    Returns:
        DocumentRequest object
    """
    # CAHMS-specific document mappings
    cahms_document_mappings = {
        "form_s": "Form S",
        "form_h": "Form H", 
        "form_a": "Form A",
        "cahms_initial": "CAHMS Initial Assessment",
        "neuro_dev_history": "Neuro Dev History",
        "formulation_document": "Formulation Document",
        "school_observation": "School Observation",
        "supporting_information": "Supporting Information"
    }
    
    # CAHMS-specific mandatory defaults
    if mandatory_uploads is None:
        mandatory_uploads = {
            "form_s": True,
            "form_h": False,
            "form_a": False,
            "cahms_initial": False,
            "neuro_dev_history": False,
            "formulation_document": False,
            "school_observation": False,
            "supporting_information": False
        }
    
    return await process_documents(
        uploaded_files=uploaded_files,
        session_id=session_id,
        document_mappings=cahms_document_mappings,
        mandatory_uploads=mandatory_uploads
    )
